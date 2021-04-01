<<<<<<< HEAD
"""
作者:Jetery
日期:2021//03//22//
"""


import docx


doc = docx.Document('demo.docx')    #demo.docx为文件名，可以按自己需求修改
def getText(filename) :
    fullText = []
    for paragraph in doc.paragraphs:
        assert isinstance(paragraph, object)
        fullText.append(paragraph.text)
    str = '\n'.join(fullText)
    return str
=======
"""
作者:Jetery
日期:2021//03//22//
"""


import docx


doc = docx.Document('demo.docx')    #demo.docx为文件名，可以按自己需求修改
def getText(filename) :
    fullText = []
    for paragraph in doc.paragraphs:
        assert isinstance(paragraph, object)
        fullText.append(paragraph.text)
    str = '\n'.join(fullText)
    return str
>>>>>>> 164e6676e9948185caa5bb9073f9092dfca9b3f3
